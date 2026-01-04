"""
Word Processor - ×¢×™×‘×•×“ ×§×‘×¦×™ Word
×¤×¨×•×™×§×˜ ×’××¨ - ×–×™×”×•×™ ××™×“×¢ ××™×©×™ ×¨×’×™×©

××•×“×•×œ ×œ×—×™×œ×•×¥ ×•× ×™×ª×•×— ×˜×§×¡×˜ ×××¡××›×™ Word
"""

from docx import Document
import io
import logging
from typing import Dict, List, Union
import os


class WordProcessor:
    """
    ××¢×‘×“ Word ×¢× ×ª××™×›×” ××œ××” ×‘-docx
    """

    def __init__(self):
        """××ª×—×•×œ ×”××¢×‘×“"""
        self.setup_logging()

    def setup_logging(self):
        """×”×’×“×¨×ª ×œ×•×’×™×"""
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def extract_text_from_word(self, word_data: Union[str, bytes],
                               filename: str = "") -> Dict:
        """
        ×—×™×œ×•×¥ ×˜×§×¡×˜ ×××¡××š Word
        """
        try:
            # ×§×¨×™××ª ×”××¡××š
            if isinstance(word_data, str):
                doc = Document(word_data)
            elif isinstance(word_data, bytes):
                doc = Document(io.BytesIO(word_data))
            else:
                raise ValueError("×¡×•×’ × ×ª×•× ×™ Word ×œ× × ×ª××š")

            self.logger.info(f"ğŸ“„ ×¢×™×‘×•×“ Word: {len(doc.paragraphs)} ×¤×¡×§××•×ª")

            # ×—×™×œ×•×¥ ×˜×§×¡×˜ ××›×œ ×”×¤×¡×§××•×ª
            paragraphs_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs_text.append(para.text)

            # ×—×™×œ×•×¥ ×˜×§×¡×˜ ××˜×‘×œ××•×ª
            tables_text = []
            for table in doc.tables:
                table_content = self._extract_table_text(table)
                if table_content:
                    tables_text.append(table_content)

            # ×—×™×œ×•×¥ ×˜×§×¡×˜ ××›×•×ª×¨×•×ª ×¢×œ×™×•× ×•×ª ×•×ª×—×ª×•× ×•×ª
            headers_footers_text = self._extract_headers_footers(doc)

            # ××™×—×•×“ ×›×œ ×”×˜×§×¡×˜
            all_text_parts = []

            if paragraphs_text:
                all_text_parts.append("=== ×ª×•×›×Ÿ ×¨××©×™ ===\n" + "\n\n".join(paragraphs_text))

            if tables_text:
                all_text_parts.append("\n\n=== ×˜×‘×œ××•×ª ===\n" + "\n\n".join(tables_text))

            if headers_footers_text:
                all_text_parts.append("\n\n=== ×›×•×ª×¨×•×ª ×¢×œ×™×•× ×•×ª/×ª×—×ª×•× ×•×ª ===\n" + headers_footers_text)

            full_text = "\n".join(all_text_parts)

            result = {
                'success': True,
                'text': full_text,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'filename': filename,
                'character_count': len(full_text),
                'word_count': len(full_text.split()) if full_text else 0,
                'has_tables': len(doc.tables) > 0,
                'has_images': self._count_images(doc) > 0,
                'image_count': self._count_images(doc)
            }

            self.logger.info(f"âœ… Word: {len(full_text)} ×ª×•×•×™×, {len(doc.paragraphs)} ×¤×¡×§××•×ª")
            return result

        except Exception as e:
            self.logger.error(f"âŒ ×©×’×™××” ×‘×¢×™×‘×•×“ Word: {e}")
            return {
                'success': False,
                'error': str(e),
                'text': "",
                'paragraphs': 0,
                'tables': 0,
                'filename': filename
            }

    def _extract_table_text(self, table) -> str:
        """×—×™×œ×•×¥ ×˜×§×¡×˜ ××˜×‘×œ×”"""
        try:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())

                if row_text:
                    table_text.append(" | ".join(row_text))

            return "\n".join(table_text)

        except Exception as e:
            self.logger.error(f"âŒ ×©×’×™××” ×‘×—×™×œ×•×¥ ×˜×‘×œ×”: {e}")
            return ""

    def _extract_headers_footers(self, doc) -> str:
        """×—×™×œ×•×¥ ×˜×§×¡×˜ ××›×•×ª×¨×•×ª ×¢×œ×™×•× ×•×ª ×•×ª×—×ª×•× ×•×ª"""
        try:
            headers_footers = []

            for section in doc.sections:
                if section.header:
                    header_text = "\n".join([para.text for para in section.header.paragraphs if para.text.strip()])
                    if header_text:
                        headers_footers.append(f"×›×•×ª×¨×ª ×¢×œ×™×•× ×”: {header_text}")

                if section.footer:
                    footer_text = "\n".join([para.text for para in section.footer.paragraphs if para.text.strip()])
                    if footer_text:
                        headers_footers.append(f"×›×•×ª×¨×ª ×ª×—×ª×•× ×”: {footer_text}")

            return "\n".join(headers_footers)

        except Exception as e:
            self.logger.error(f"âŒ ×©×’×™××” ×‘×—×™×œ×•×¥ ×›×•×ª×¨×•×ª: {e}")
            return ""

    def _count_images(self, doc) -> int:
        """×¡×¤×™×¨×ª ×ª××•× ×•×ª ×‘××¡××š"""
        try:
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            return image_count
        except:
            return 0

    def get_word_info(self, word_data: Union[str, bytes]) -> Dict:
        """×§×‘×œ×ª ××™×“×¢ ×¢×œ ××¡××š Word"""
        try:
            if isinstance(word_data, str):
                doc = Document(word_data)
                file_size = os.path.getsize(word_data)
            else:
                doc = Document(io.BytesIO(word_data))
                file_size = len(word_data)

            core_props = doc.core_properties

            info = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
                'images': self._count_images(doc),
                'file_size': file_size,
                'title': core_props.title if core_props.title else '',
                'author': core_props.author if core_props.author else '',
                'subject': core_props.subject if core_props.subject else '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }

            return info

        except Exception as e:
            self.logger.error(f"âŒ ×©×’×™××” ×‘×§×‘×œ×ª ××™×“×¢ Word: {e}")
            return {}


def is_word_file(filename: str) -> bool:
    """×‘×“×™×§×” ×× ×”×§×•×‘×¥ ×”×•× Word"""
    if not filename:
        return False
    return filename.lower().endswith(('.docx', '.doc'))


if __name__ == "__main__":
    print("ğŸ“„ ×‘×“×™×§×ª ××¢×‘×“ Word")
    print("=" * 30)
    processor = WordProcessor()
    print("âœ… ××¢×‘×“ Word ××•×›×Ÿ ×œ×©×™××•×©!")