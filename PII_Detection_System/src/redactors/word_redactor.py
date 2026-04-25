"""
Word Redactor - השחרת קבצי Word
פרויקט גמר - זיהוי מידע אישי רגיש

מודול להשחרה אקטיבית של PII מקבצי Word
"""

import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
import io
import logging
import random
from typing import List, Union, Optional

class WordRedactor:
    """
    מחלקה האחראית על השחרת מידע רגיש מקבצי Word.
    המחלקה מקבלת קובץ ורשימת טקסטים להסרה, משנה את הטקסט באופן בלתי הפיך ל-0 ול-1
    ומייצרת קובץ Word חדש ובטוח.
    """

    def __init__(self):
        self._setup_logging()

    def _setup_logging(self):
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def _mask_text_binary(self, original_text: str) -> str:
        """מחליף כל תו בטקסט הרגיש ב-0 או 1 אקראיים כדי למנוע שחזור"""
        return "".join(random.choice(['0', '1']) for _ in original_text)

    def _redact_paragraph(self, para, pii_texts: List[str]) -> int:
        redact_count = 0
        
        # שלב 1: ניסיון השחרה ברמת ה-Run (שומר על עיצוב מקורי)
        for run in para.runs:
            if not run.text: 
                continue
            modified = False
            new_text = run.text
            for pii in pii_texts:
                if pii and pii in new_text:
                    masked = self._mask_text_binary(pii)
                    new_text = new_text.replace(pii, masked)
                    modified = True
                    redact_count += 1
            if modified:
                run.text = new_text
                # צביעת הרקע בשחור והטקסט בלבן
                run.font.highlight_color = WD_COLOR_INDEX.BLACK
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        # שלב 2: בדיקה אם נשאר מידע רגיש שפוצל בין כמה Runs
        remaining_pii = False
        para_text = para.text
        for pii in pii_texts:
            if pii and pii in para_text:
                remaining_pii = True
                break
                
        if remaining_pii:
            # Fallback: דריסת הפסקה כולה כדי להבטיח מחיקה (על חשבון עיצוב)
            new_text = para.text
            for pii in pii_texts:
                if pii and pii in new_text:
                    masked = self._mask_text_binary(pii)
                    new_text = new_text.replace(pii, masked)
                    redact_count += 1
            para.clear()
            run = para.add_run(new_text)
            
        return redact_count

    def redact_word(self, word_data: Union[str, bytes], pii_texts: List[str], output_path: Optional[str] = None) -> Union[bytes, str, bool]:
        """
        השחרת נתונים רגישים מתוך קובץ Word.
        - word_data: נתיב לקובץ מקור או נתוני bytes
        - pii_texts: רשימה של מחרוזות (הטקסט של המידע הרגיש) שיש להשחיר
        - output_path: נתיב לשמירת הקובץ. אם None, יוחזרו bytes.
        """
        try:
            if isinstance(word_data, str):
                doc = docx.Document(word_data)
            else:
                doc = docx.Document(io.BytesIO(word_data))
                
            self.logger.info(f"🔒 מתחיל השחרת קובץ Word. מספר מחרוזות PII להשחרה: {len(pii_texts)}")
            total_redact_count = 0

            # 1. מעבר על פסקאות רגילות
            for para in doc.paragraphs:
                total_redact_count += self._redact_paragraph(para, pii_texts)

            # 2. מעבר על פסקאות בתוך טבלאות
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            total_redact_count += self._redact_paragraph(para, pii_texts)

            self.logger.info(f"✅ סיום השחרה. {total_redact_count} חלקי מידע הופכו למחרוזת בינארית (0 ו-1).")

            # שמירת התוצאה - החזרת הקובץ למשתמש להורדה
            if output_path is not None:
                doc.save(output_path)
                return output_path
            else:
                output = io.BytesIO()
                doc.save(output)
                return output.getvalue()
                
        except Exception as e:
            self.logger.error(f"❌ שגיאה בהשחרת Word: {e}")
            return False