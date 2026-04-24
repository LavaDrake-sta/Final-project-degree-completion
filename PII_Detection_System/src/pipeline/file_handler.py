import os
import io
import mimetypes
import pdfplumber
import docx
import openpyxl
from PIL import Image
from typing import Union, Dict, Any

from .ocr_processor import OCRProcessor

class FileHandler:
    """
    Handles extraction of text from various file formats (PDF, DOCX, XLSX, Images).
    """

    def __init__(self):
        self.ocr_processor = OCRProcessor()

    def detect_file_type(self, file_path: str = None, file_bytes: bytes = None, filename: str = None) -> str:
        """Detect the type of the file based on extension or mime type."""
        name = file_path if file_path else filename
        if not name:
            return "unknown"
        
        name = name.lower()
        if name.endswith('.pdf'):
            return "pdf"
        elif name.endswith('.docx'):
            return "docx"
        elif name.endswith('.xlsx'):
            return "xlsx"
        elif name.endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
            return "image"
        else:
            mime_type, _ = mimetypes.guess_type(name)
            if mime_type and mime_type.startswith('image/'):
                return "image"
            return "unknown"

    def process_file(self, file_path: str = None, file_bytes: bytes = None, filename: str = None) -> Dict[str, Any]:
        """Process the file and extract text and metadata."""
        file_type = self.detect_file_type(file_path, file_bytes, filename)
        
        if file_type == "pdf":
            return self._process_pdf(file_path, file_bytes)
        elif file_type == "docx":
            return self._process_docx(file_path, file_bytes)
        elif file_type == "xlsx":
            return self._process_xlsx(file_path, file_bytes)
        elif file_type == "image":
            return self._process_image(file_path, file_bytes)
        else:
            return {"success": False, "error": f"Unsupported file type: {file_type}", "text": ""}

    def _process_pdf(self, file_path: str = None, file_bytes: bytes = None) -> Dict[str, Any]:
        extracted_text = []
        is_scanned = False
        
        try:
            if file_path:
                pdf = pdfplumber.open(file_path)
            else:
                pdf = pdfplumber.open(io.BytesIO(file_bytes))
                
            with pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and len(text.strip()) > 10:
                        extracted_text.append(text)
                    else:
                        # Convert page to image and run OCR
                        im = page.to_image(resolution=300)
                        pil_image = im.original
                        ocr_text = self.ocr_processor.extract_from_image_obj(pil_image)
                        extracted_text.append(ocr_text)
                        is_scanned = True
                        
            full_text = "\n".join(extracted_text)
            return {
                "success": True, 
                "text": full_text, 
                "is_scanned": is_scanned,
                "file_type": "pdf"
            }
        except Exception as e:
            return {"success": False, "error": str(e), "text": ""}

    def _process_docx(self, file_path: str = None, file_bytes: bytes = None) -> Dict[str, Any]:
        try:
            if file_path:
                doc = docx.Document(file_path)
            else:
                doc = docx.Document(io.BytesIO(file_bytes))
                
            text = "\n".join([para.text for para in doc.paragraphs])
            return {"success": True, "text": text, "file_type": "docx"}
        except Exception as e:
            return {"success": False, "error": str(e), "text": ""}

    def _process_xlsx(self, file_path: str = None, file_bytes: bytes = None) -> Dict[str, Any]:
        try:
            if file_path:
                wb = openpyxl.load_workbook(file_path, data_only=True)
            else:
                wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
                
            extracted_text = []
            for sheet in wb.worksheets:
                extracted_text.append(f"--- Sheet: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) for cell in row if cell is not None])
                    if row_text.strip():
                        extracted_text.append(row_text)
                        
            return {"success": True, "text": "\n".join(extracted_text), "file_type": "xlsx"}
        except Exception as e:
            return {"success": False, "error": str(e), "text": ""}

    def _process_image(self, file_path: str = None, file_bytes: bytes = None) -> Dict[str, Any]:
        try:
            if file_path:
                text = self.ocr_processor.extract_from_path(file_path)
            else:
                text = self.ocr_processor.extract_from_bytes(file_bytes)
                
            return {"success": True, "text": text, "file_type": "image"}
        except Exception as e:
            return {"success": False, "error": str(e), "text": ""}
