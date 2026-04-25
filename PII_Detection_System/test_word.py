import sys
import os
import docx

# Add src to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), 'src')))

from redactors.word_redactor import WordRedactor

def test_word_redactor():
    print("Testing Word Redactor...")
    
    # Create a dummy Word file
    doc = docx.Document()
    
    doc.add_heading('מסמך בדיקה', 0)
    
    p = doc.add_paragraph('שם הלקוח הוא ')
    p.add_run('ישראל ישראלי').bold = True
    p.add_run(' והוא גר ב')
    p.add_run('תל אביב').italic = True
    
    doc.add_paragraph('ת.ז של הלקוח: 123456789 (סודי בהחלט)')
    
    # Add a table to check table redaction
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'טלפון נייד'
    hdr_cells[1].text = '052-1234567'
    
    dummy_path = "dummy_test.docx"
    doc.save(dummy_path)
    
    # Initialize redactor
    redactor = WordRedactor()
    
    # The PII strings we pretend the system detected
    pii_texts = ["123456789", "052-1234567", "ישראל ישראלי"]
    
    output_path = "dummy_test_redacted.docx"
    
    result = redactor.redact_word(dummy_path, pii_texts, output_path)
    
    if result:
        print(f"Redaction successful! Saved to {result}")
    else:
        print("Redaction failed!")
        
    # Clean up (optional)
    if os.path.exists(dummy_path):
        os.remove(dummy_path)

if __name__ == "__main__":
    test_word_redactor()
