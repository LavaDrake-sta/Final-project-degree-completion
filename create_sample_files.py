"""
Sample File Generator
Creates sample files for testing the PII detection system
"""

from pathlib import Path
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PIL import Image, ImageDraw, ImageFont

# Sample text with PII (Hebrew and English)
SAMPLE_TEXT_1 = """
פרטים אישיים של הלקוח

שם: דוד כהן
תעודת זהות: 123456782
טלפון: 050-1234567
דוא"ל: david.cohen@example.com
כתובת: רחוב הרצל 45, תל אביב

פרטי תשלום:
חשבון בנק: 12-345678-01
כרטיס אשראי: 4580-1234-5678-9012
"""

SAMPLE_TEXT_2 = """
Customer Information Form

Name: Sarah Levi
ID Number: 234567893
Phone: +972-52-9876543
Email: sarah.levi@company.co.il
Address: Dizengoff Street 123, Tel Aviv

The customer requested information about our services.
Contact person: Michael Green (michael@company.com)
Organization: Tech Solutions Ltd.
Location: Jerusalem, Israel
"""

SAMPLE_TEXT_3 = """
דוח פגישה

משתתפים:
1. יוסי אברהם - יו"ר (yossi.a@example.org)
2. רחל ברוך - מנכ"לית
3. משה דוד - מנהל כספים

נושאים שנדונו:
- עדכון פרטי לקוחות במערכת
- שדרוג אבטחת מידע
- טלפונים ליצירת קשר: 03-1234567, 054-9876543

כתובת המשרד: שדרות רוטשילד 1, תל אביב
"""


def create_sample_docx(filename: str, text: str, output_dir: str = "data/input"):
    """Create a sample DOCX file"""
    output_path = Path(output_dir) / filename

    doc = Document()
    doc.add_heading('מסמך לדוגמה', 0)

    for paragraph in text.strip().split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)

    doc.save(output_path)
    print(f"✓ נוצר: {output_path}")


def create_sample_txt(filename: str, text: str, output_dir: str = "data/input"):
    """Create a sample text file"""
    output_path = Path(output_dir) / filename

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

    print(f"✓ נוצר: {output_path}")


def create_sample_image_simple(filename: str, text: str, output_dir: str = "data/input"):
    """Create a simple sample image with text (without special fonts)"""
    output_path = Path(output_dir) / filename

    # Create image
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)

    # Use default font (will work but may not support Hebrew perfectly)
    try:
        # Try to use a larger default font
        font = ImageFont.load_default()
    except:
        font = None

    # Draw text
    y_position = 50
    for line in text.strip().split('\n'):
        if line.strip():
            draw.text((50, y_position), line, fill='black', font=font)
            y_position += 30

    img.save(output_path)
    print(f"✓ נוצר: {output_path}")


def main():
    """Generate sample files"""
    print("=" * 60)
    print("יצירת קבצי דוגמה לבדיקת המערכת")
    print("=" * 60)
    print()

    # Ensure output directory exists
    output_dir = Path("data/input")
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        # Create DOCX files
        print("יוצר קבצי Word...")
        create_sample_docx("sample_hebrew.docx", SAMPLE_TEXT_1)
        create_sample_docx("sample_english.docx", SAMPLE_TEXT_2)

        # Create text files
        print("\nיוצר קבצי טקסט...")
        create_sample_txt("sample_meeting.txt", SAMPLE_TEXT_3)

        # Create image files (simple, without Hebrew font)
        print("\nיוצר קבצי תמונה...")
        create_sample_image_simple("sample_image.png", SAMPLE_TEXT_2)

        print()
        print("=" * 60)
        print("✨ כל הקבצים נוצרו בהצלחה!")
        print(f"📁 מיקום: {output_dir.absolute()}")
        print()
        print("כעת ניתן להריץ את המערכת:")
        print("  python -m src.main")
        print("=" * 60)

    except ImportError as e:
        print(f"\n⚠️  חסרות ספריות: {e}")
        print("התקן את הספריות הנדרשות:")
        print("  pip install python-docx pillow")
    except Exception as e:
        print(f"\n❌ שגיאה: {e}")


if __name__ == "__main__":
    main()
